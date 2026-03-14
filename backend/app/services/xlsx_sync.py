import io
from typing import List
import pandas as pd
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class XlsxSyncService:
    def parse_excel_to_items(self, xlsx_content: bytes) -> List[dict]:
        df = pd.read_excel(io.BytesIO(xlsx_content))
        return df.fillna("N/A").to_dict(orient='records')

    def generate_docx_report(self, items: List[dict]) -> bytes:
        doc = Document()
        
        # Title
        title = doc.add_heading('Operational Governance Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = RGBColor(0xFF, 0x62, 0x00)

        doc.add_paragraph("Deterministic Transformation Engine (Standard Compliant)", style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # Group by category
        grouped = {}
        for item in items:
            cat = item.get('Category', 'General')
            if cat not in grouped: grouped[cat] = []
            grouped[cat].append(item)

        for cat, cat_items in grouped.items():
            heading = doc.add_heading(cat, level=1)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0xFF, 0x62, 0x00)
            
            for i, item in enumerate(cat_items):
                p = doc.add_paragraph()
                p.add_run(f"Finding {i+1}: ").bold = True
                p.add_run(str(item.get('Finding', 'N/A')))
                
                p = doc.add_paragraph()
                p.add_run(f"Action: ").bold = True
                p.add_run(str(item.get('Action Item', 'N/A')))
                
                # Metadata line
                meta = doc.add_paragraph()
                meta_str = f"Owner: {item.get('Owner', 'N/A')} | Priority: {item.get('Priority', 'N/A')} | Status: {item.get('Status', 'N/A')}"
                meta_run = meta.add_run(meta_str)
                meta_run.font.size = RGBColor(0x66, 0x66, 0x66) # Small Gray font (simulated via manual Pt size if needed)
                meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                doc.add_paragraph() # Spacer

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

xlsx_sync_service = XlsxSyncService()
